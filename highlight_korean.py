#!/usr/bin/env python3
"""
Highlights glossary terms in the Korean document
Creates a new document with the same content, with terms highlighted in yellow
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import re
from typing import List, Tuple, Dict
from collections import Counter


# -----------------------------
# Korean particles to ignore
# -----------------------------
KO_PARTICLES = {"는","은","이","가","을","를","에","도","와","과","한","의"}


# -----------------------------
# Load Glossary
# -----------------------------
def load_glossary(file_path: str) -> List[Tuple[str, str]]:
    df = pd.read_excel(file_path)

    ko_terms = df.iloc[:, 5].astype(str)
    en_terms = df.iloc[:, 6].astype(str)

    lingo = dict(zip(ko_terms, en_terms))

    filtered = {
        k.strip(): v.strip()
        for k, v in lingo.items()
        if k not in KO_PARTICLES and k != "nan" and v != "nan" and k.strip()
    }

    # 🔥 IMPORTANT: longest first
    sorted_terms = sorted(filtered.keys(), key=len, reverse=True)

    return [(ko, filtered[ko]) for ko in sorted_terms]


def build_pattern(glossary: List[Tuple[str, str]]):

    escaped_terms = [re.escape(ko) for ko, _ in glossary]

    pattern = re.compile(
        r'(?<![가-힣])(' + '|'.join(escaped_terms) + r')(?![가-힣])'
    )

    lookup = {ko: (ko, en) for ko, en in glossary}

    return pattern, lookup


# -----------------------------
# Highlight Styling
# -----------------------------
def add_highlight(run, color="FFFF00"):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    run._element.get_or_add_rPr().append(shd)


# -----------------------------
# Highlight Paragraph
# -----------------------------
def highlight_paragraph(para, pattern, lookup, terms_found):

    if not para.runs:
        return

    full_text = "".join(run.text for run in para.runs)

    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    # Because regex already prefers leftmost-longest,
    # we don't need manual overlap removal

    # Track matches
    for m in matches:
        key = m.group(0)
        if key in lookup:
            ko, en = lookup[key]
            terms_found.append({
                "korean": ko,
                "english": en,
                "found_text": key
            })

    # Rebuild paragraph
    para.clear()

    last = 0

    for m in matches:
        start, end = m.start(), m.end()

        # Add normal text
        if start > last:
            para.add_run(full_text[last:start])

        # Add highlighted text
        run = para.add_run(full_text[start:end])
        add_highlight(run)

        last = end

    # Tail text
    if last < len(full_text):
        para.add_run(full_text[last:])


# -----------------------------
# Process Document
# -----------------------------
def process_document(doc, pattern, lookup):

    terms_found = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            highlight_paragraph(para, pattern, lookup, terms_found)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        highlight_paragraph(para, pattern, lookup, terms_found)

    return terms_found


# -----------------------------
# Main Function
# -----------------------------
def create_highlighted_korean_doc(ko_doc_path, glossary_path, output_path):

    print(f"Loading document: {ko_doc_path}")
    doc = Document(ko_doc_path)

    print("Loading glossary...")
    glossary = load_glossary(glossary_path)
    print(f"Loaded {len(glossary)} terms")

    print("Building regex...")
    pattern, lookup = build_pattern(glossary)

    print("Processing document...")
    terms_found = process_document(doc, pattern, lookup)

    doc.save(output_path)

    print(f"\nSaved → {output_path}")
    print(f"Total matches: {len(terms_found)}")

    return terms_found


# -----------------------------
# Run
# -----------------------------
if __name__ == "__main__":

    terms_found = create_highlighted_korean_doc(
        ko_doc_path="ko/KO-Test.docx",
        glossary_path="glossary/L2M-OOG-Lingo-0313.xlsx",
        output_path="output/KO-Test_Highlighted.docx"
    )

    print("\n" + "=" * 80)
    print(f"Total occurrences: {len(terms_found)}")
    print(f"Unique terms: {len(set(t['korean'] for t in terms_found))}")

    term_counts = Counter(t['korean'] for t in terms_found)

    print("\nTop 10 most common terms:")
    for korean, count in term_counts.most_common(10):
        english = next(t['english'] for t in terms_found if t['korean'] == korean)
        print(f"{count:3d}x {korean:30} -> {english}")