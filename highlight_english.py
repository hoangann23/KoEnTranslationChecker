#!/usr/bin/env python3
"""
Highlights glossary terms in the English document
Creates a new document with the same content, with terms highlighted in yellow
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import re
from typing import List, Tuple, Dict
from collections import Counter


# -----------------------------
# Load Glossary
# -----------------------------
def load_glossary(file_path: str) -> List[Tuple[str, str]]:
    df = pd.read_excel(file_path)

    ko_terms = df.iloc[:, 5].astype(str)
    en_terms = df.iloc[:, 6].astype(str)

    lingo = dict(zip(en_terms, ko_terms))

    filtered = {
        k.strip(): v.strip()
        for k, v in lingo.items()
        if k != "nan" and v != "nan" and k.strip()
    }

    # Sort longest first (important!)
    sorted_terms = sorted(filtered.keys(), key=len, reverse=True)

    return [(en, filtered[en]) for en in sorted_terms]


# -----------------------------
# Build Regex Pattern (ONE TIME)
# -----------------------------
def build_pattern(glossary: List[Tuple[str, str]]):
    escaped = [re.escape(en) for en, _ in glossary]

    pattern = re.compile(
        r'(?<!\w)(' + '|'.join(escaped) + r')(?!\w)',
        re.IGNORECASE
    )

    # Map lowercase → original glossary entry
    lookup = {en.lower(): (en, ko) for en, ko in glossary}

    return pattern, lookup


# -----------------------------
# Highlight Styling
# -----------------------------
def add_highlight(run, color="FFFF00"):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    run._element.get_or_add_rPr().append(shd)


# -----------------------------
# Highlight Paragraph (SAFE VERSION)
# -----------------------------
def highlight_paragraph(para, pattern, lookup, terms_found):

    original_runs = para.runs
    if not original_runs:
        return

    full_text = "".join(run.text for run in original_runs)

    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    # Track matches
    for m in matches:
        key = m.group(0).lower()
        if key in lookup:
            en, ko = lookup[key]
            terms_found.append({
                "english": en,
                "korean": ko,
                "found_text": m.group(0)
            })

    # Clear paragraph
    para.clear()

    last = 0

    for m in matches:
        start, end = m.start(), m.end()

        # Normal text
        if start > last:
            para.add_run(full_text[last:start])

        # Highlighted text
        run = para.add_run(full_text[start:end])
        add_highlight(run)

        last = end

    # Tail text
    if last < len(full_text):
        para.add_run(full_text[last:])


# -----------------------------
# Process Document
# -----------------------------
def process_document(doc, pattern, lookup):
    terms_found = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            highlight_paragraph(para, pattern, lookup, terms_found)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        highlight_paragraph(para, pattern, lookup, terms_found)

    return terms_found


# -----------------------------
# Main Function
# -----------------------------
def create_highlighted_doc(en_doc_path, glossary_path, output_path):
    print(f"Loading document: {en_doc_path}")
    doc = Document(en_doc_path)

    print("Loading glossary...")
    glossary = load_glossary(glossary_path)

    print(f"Loaded {len(glossary)} terms")

    print("Building regex...")
    pattern, lookup = build_pattern(glossary)

    print("Processing document...")
    terms_found = process_document(doc, pattern, lookup)

    doc.save(output_path)

    print(f"\nSaved → {output_path}")
    print(f"Total matches: {len(terms_found)}")

    return terms_found


# -----------------------------
# Run
# -----------------------------
if __name__ == "__main__":

    terms_found = create_highlighted_doc(
        en_doc_path="en/EN-Test.docx",
        glossary_path="glossary/L2M-OOG-Lingo-0313.xlsx",
        output_path="output/EN-Test_Highlighted.docx"
    )

    print("\n" + "=" * 80)
    print(f"Total occurrences: {len(terms_found)}")
    print(f"Unique terms: {len(set(t['english'] for t in terms_found))}")

    term_counts = Counter(t['english'] for t in terms_found)

    print("\nTop 10 most common terms:")
    for english, count in term_counts.most_common(10):
        korean = next(t['korean'] for t in terms_found if t['english'] == english)
        print(f"{count:3d}x {english:40} <- {korean}")